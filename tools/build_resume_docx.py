#!/usr/bin/env python3
from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import UTC, datetime
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "resumes/current/Jordan_Example_Resume.md"
OUT = ROOT / "resumes/current/Jordan_Example_Resume.docx"
STYLE_PATH = ROOT / "config/resume-style.yaml"

DEFAULT_STYLE = {
    "font": {
        "family": "Calibri",
        "name_size": 18,
        "contact_size": 8.8,
        "normal_size": 8.8,
        "summary_size": 8.6,
        "body_size": 8.5,
        "skill_label_size": 8.45,
        "skill_text_size": 8.35,
        "bullet_size": 8.45,
        "section_size": 9.4,
        "role_size": 9.1,
        "subrole_size": 8.75,
    },
    "color": {"accent": "1F4E79", "text": "232323"},
    "margins": {"top": 0.42, "bottom": 0.42, "left": 0.55, "right": 0.55},
    "spacing": {
        "name_after": 1,
        "contact_after": 5,
        "section_before": 2.0,
        "section_after": 0.8,
        "role_before": 1.5,
        "subrole_before": 0.2,
        "role_after": 0.2,
        "paragraph_after": 1.2,
        "skill_after": 0.7,
        "bullet_after": 1.15,
    },
    "line_spacing": {"paragraph": 1.0, "skill": 0.98, "bullet": 0.98},
    "indent": {"bullet_left": 0.18, "bullet_first_line": -0.11},
}


@dataclass
class ParsedResume:
    name: str
    contact: str
    body: list[str]


def deep_merge(base: dict, override: dict) -> dict:
    merged = dict(base)
    for key, value in override.items():
        if isinstance(value, dict) and isinstance(merged.get(key), dict):
            merged[key] = deep_merge(merged[key], value)
        else:
            merged[key] = value
    return merged


def load_style() -> dict:
    if not STYLE_PATH.exists():
        return DEFAULT_STYLE
    loaded = yaml.safe_load(STYLE_PATH.read_text(encoding="utf-8")) or {}
    if not isinstance(loaded, dict):
        raise ValueError(f"Style config must be a mapping: {STYLE_PATH}")
    return deep_merge(DEFAULT_STYLE, loaded)


STYLE = load_style()


def style(path: str):
    value = STYLE
    for part in path.split("."):
        value = value[part]
    return value


def rgb(hex_value: str) -> RGBColor:
    clean = hex_value.strip().lstrip("#")
    return RGBColor(int(clean[0:2], 16), int(clean[2:4], 16), int(clean[4:6], 16))


ACCENT = rgb(style("color.accent"))
TEXT = rgb(style("color.text"))


def set_margins(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(style("margins.top"))
    section.bottom_margin = Inches(style("margins.bottom"))
    section.left_margin = Inches(style("margins.left"))
    section.right_margin = Inches(style("margins.right"))


def set_document_defaults(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    family = style("font.family")
    normal.font.name = family
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), family)
    normal.font.size = Pt(style("font.normal_size"))
    normal.font.color.rgb = TEXT

    for style_name in ["List Bullet", "List Bullet 2"]:
        list_style = styles[style_name]
        list_style.font.name = family
        list_style._element.rPr.rFonts.set(qn("w:eastAsia"), family)
        list_style.font.size = Pt(style("font.bullet_size"))
        list_style.paragraph_format.space_after = Pt(style("spacing.bullet_after"))
        list_style.paragraph_format.line_spacing = style("line_spacing.bullet")


def set_spacing(paragraph, before: float = 0, after: float = 2, line: float = 1.0) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_bottom_border(paragraph, color: str = "4F81BD", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def add_hyperlink_run(paragraph, text: str, url: str, size: float = 8.95) -> None:
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), str(style("color.accent")))
    r_pr.append(color)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(underline)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(int(size * 2)))
    r_pr.append(sz)
    text_el = OxmlElement("w:t")
    text_el.text = text
    new_run.append(r_pr)
    new_run.append(text_el)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline_markdown(paragraph, text: str, size: float = 8.95) -> None:
    pos = 0
    pattern = re.compile(r"(\[([^\]]+)\]\(([^)]+)\)|\*\*([^*]+)\*\*)")
    for match in pattern.finditer(text):
        if match.start() > pos:
            run = paragraph.add_run(text[pos : match.start()])
            run.font.size = Pt(size)
        if match.group(2) is not None:
            add_hyperlink_run(paragraph, match.group(2), match.group(3), size=size)
        else:
            run = paragraph.add_run(match.group(4))
            run.bold = True
            run.font.size = Pt(size)
        pos = match.end()
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.size = Pt(size)


def normalize_url(value: str) -> str:
    if value.startswith("http") or value.startswith("mailto:"):
        return value
    if "@" in value:
        return f"mailto:{value}"
    return value


def add_contact_token(paragraph, token: str) -> None:
    token = token.strip()
    if not token:
        return

    label, _, value = token.partition(":")
    value = value.strip()
    if label == "Email":
        add_hyperlink_run(paragraph, value, f"mailto:{value}", size=style("font.contact_size"))
    elif label in {"GitHub", "LinkedIn", "Portfolio", "Website"}:
        add_hyperlink_run(paragraph, label, normalize_url(value), size=style("font.contact_size"))
    elif value:
        run = paragraph.add_run(value)
        run.font.size = Pt(style("font.contact_size"))
    else:
        run = paragraph.add_run(token)
        run.font.size = Pt(style("font.contact_size"))


def add_name_and_contact(doc: Document, parsed: ParsedResume) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=style("spacing.name_after"))
    run = p.add_run(parsed.name)
    run.bold = True
    run.font.size = Pt(style("font.name_size"))
    run.font.color.rgb = ACCENT

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_spacing(p, after=style("spacing.contact_after"))
    contact = parsed.contact
    for index, token in enumerate(contact.split("|")):
        if index:
            sep = p.add_run(" | ")
            sep.font.size = Pt(style("font.contact_size"))
        add_contact_token(p, token)


def add_section(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    set_spacing(p, before=style("spacing.section_before"), after=style("spacing.section_after"))
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(style("font.section_size"))
    run.font.color.rgb = ACCENT
    add_bottom_border(p)


def add_role_header(doc: Document, text: str, level: int) -> None:
    p = doc.add_paragraph()
    set_spacing(
        p,
        before=style("spacing.role_before") if level == 3 else style("spacing.subrole_before"),
        after=style("spacing.role_after"),
    )
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(style("font.role_size") if level == 3 else style("font.subrole_size"))
    run.font.color.rgb = ACCENT if level == 3 else TEXT


def add_paragraph_text(
    doc: Document,
    text: str,
    size: float | None = None,
    after: float | None = None,
) -> None:
    p = doc.add_paragraph()
    set_spacing(
        p,
        after=style("spacing.paragraph_after") if after is None else after,
        line=style("line_spacing.paragraph"),
    )
    add_inline_markdown(p, text, size=style("font.body_size") if size is None else size)


def add_skill(doc: Document, line: str) -> None:
    match = re.match(r"\*\*([^:]+):\*\*\s*(.*)", line)
    if not match:
        add_paragraph_text(doc, line)
        return
    p = doc.add_paragraph()
    set_spacing(p, after=style("spacing.skill_after"), line=style("line_spacing.skill"))
    label = p.add_run(f"{match.group(1)}: ")
    label.bold = True
    label.font.size = Pt(style("font.skill_label_size"))
    label.font.color.rgb = ACCENT
    add_inline_markdown(p, match.group(2), size=style("font.skill_text_size"))


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_spacing(p, after=style("spacing.bullet_after"), line=style("line_spacing.bullet"))
    p.paragraph_format.left_indent = Inches(style("indent.bullet_left"))
    p.paragraph_format.first_line_indent = Inches(style("indent.bullet_first_line"))
    bullet = p.add_run("- ")
    bullet.font.size = Pt(style("font.bullet_size"))
    add_inline_markdown(p, text, size=style("font.bullet_size"))


def parse_resume(source: Path) -> ParsedResume:
    if not source.exists():
        raise FileNotFoundError(source)
    lines = source.read_text(encoding="utf-8").splitlines()
    non_empty = [line for line in lines if line.strip()]
    if not non_empty or not non_empty[0].startswith("# "):
        raise ValueError("Resume markdown must start with '# Name'.")
    name = non_empty[0][2:].strip()
    contact = next((line for line in non_empty[1:] if line.startswith("Phone: ")), "")
    body_start = lines.index(contact) + 1 if contact else 1
    return ParsedResume(name=name, contact=contact, body=lines[body_start:])


def build() -> None:
    parsed = parse_resume(SOURCE)
    doc = Document()
    now = datetime.now(UTC).replace(tzinfo=None)
    doc.core_properties.author = parsed.name
    doc.core_properties.last_modified_by = parsed.name
    doc.core_properties.title = f"{parsed.name} Resume"
    doc.core_properties.created = now
    doc.core_properties.modified = now
    set_margins(doc)
    set_document_defaults(doc)
    add_name_and_contact(doc, parsed)

    in_skills = False
    for raw in parsed.body:
        line = raw.strip()
        if not line:
            continue
        if line.startswith("## "):
            title = line[3:].strip()
            in_skills = title == "Skills"
            add_section(doc, title)
        elif line.startswith("### "):
            add_role_header(doc, line[4:].strip(), level=3)
        elif line.startswith("#### "):
            add_role_header(doc, line[5:].strip(), level=4)
        elif line.startswith("- "):
            add_bullet(doc, line[2:].strip())
        elif in_skills and line.startswith("**"):
            add_skill(doc, line)
        else:
            add_paragraph_text(
                doc,
                line,
                size=style("font.summary_size") if "Engineer and full-stack" in line else style("font.body_size"),
            )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
